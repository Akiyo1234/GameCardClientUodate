# -*- coding: utf-8 -*-
"""สร้างเอกสาร Word (.docx) จากแผนพัฒนา 4 ระบบ — ใช้ฟอนต์ไทย TH Sarabun New"""
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THAI_FONT = "TH Sarabun New"
MONO_FONT = "Consolas"

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # navy
ACCENT  = RGBColor(0x2E, 0x74, 0xB5)   # blue
MUTED   = RGBColor(0x60, 0x60, 0x60)
GREEN   = RGBColor(0x1B, 0x7A, 0x34)
RED     = RGBColor(0xB0, 0x1E, 0x1E)

doc = Document()

# ---- page A4 + margins ----
sec = doc.sections[0]
sec.page_height = Mm(297); sec.page_width = Mm(210)
sec.top_margin = Mm(20); sec.bottom_margin = Mm(20)
sec.left_margin = Mm(22); sec.right_margin = Mm(20)

def _apply_font(run, name=THAI_FONT, size=15, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)

def _shade(cell_or_para, hex_fill):
    """ใส่สีพื้นหลังให้ cell หรือ paragraph"""
    el = cell_or_para._tc if hasattr(cell_or_para, '_tc') else cell_or_para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hex_fill)
    el.append(shd)

def title(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(p.add_run(text), size=26, bold=True, color=PRIMARY)
    p.paragraph_format.space_after = Pt(2)

def subtitle(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(p.add_run(text), size=14, color=MUTED)
    p.paragraph_format.space_after = Pt(10)

def h1(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    _apply_font(p.add_run(text), size=19, bold=True, color=PRIMARY)
    # เส้นใต้หัวข้อ
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), '2E74B5')
    pbdr.append(bottom); pPr.append(pbdr)

def h2(text, color=ACCENT):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    _apply_font(p.add_run(text), size=16, bold=True, color=color)

def para(text, size=15, color=None, italic=False, space=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    _apply_font(p.add_run(text), size=size, color=color, italic=italic)
    return p

def bullet(text, level=0, check=False):
    p = doc.add_paragraph(style='List Bullet' if not check else None)
    p.paragraph_format.left_indent = Mm(8 + level*7)
    p.paragraph_format.space_after = Pt(1)
    if check:
        _apply_font(p.add_run("☐  "), size=15)
    _apply_font(p.add_run(text), size=15)
    return p

def code_block(lines):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Mm(4)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    _shade(p, 'F0F3F7')
    for i, ln in enumerate(lines):
        r = p.add_run(("" if i == 0 else "\n") + ln)
        _apply_font(r, name=MONO_FONT, size=11)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; _shade(c, '1F4E79')
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        _apply_font(c.paragraphs[0].add_run(htext), size=14, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            _apply_font(cells[i].paragraphs[0].add_run(val), size=13.5)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ================= CONTENT =================
title("แผนพัฒนา 4 ระบบ ภายใน 1 เดือน")
subtitle("GameCardClient · Unity 6 + Photon Fusion (Shared Mode) + Supabase · จัดทำ 9 มิ.ย. 2026")

h1("ภาพรวมและ Scope ที่ตกลง")
para("เอกสารนี้สรุปงาน 4 ข้อที่ต้องแก้ก่อนสอบ พร้อมสถานะปัจจุบันของโค้ด (จากการอ่านจริง) วิธีทำ "
     "ลำดับงาน และจุดเสี่ยง  ระดับของแต่ละข้อยืนยันกับเจ้าของโปรเจคแล้ว")
table(
    ["ข้อ", "สิ่งที่ต้องทำ", "ระดับที่เลือก", "ความยาก"],
    [
        ["1. Reconnect", "ผู้เล่นหลุดกลับเข้าแมตช์เดิม", "แบบเกมออนไลน์ทั่วไป — ทวง seat+state รองรับ app ปิด/เปิด", "ยากสุด"],
        ["2. Tutorial", "สอนเล่นแบบสัมผัสจริง", "Coach-mark ทับเกมจริง (เล่นกับบอท)", "กลาง"],
        ["3. Log → DB", "เก็บ log การเล่นลงฐานข้อมูล", "analytics เกม + error (schema generic)", "กลาง"],
        ["4. Quiz UI", "เพิ่มคำถามผ่าน UI ลง DB", "ฟอร์มกรอกทีละข้อในหน้าเว็บ", "ง่ายสุด"],
    ],
    widths=[26, 42, 92, 22],
)

# ---------- ข้อ 4 ----------
h1("ข้อ 4 — Quiz: ฟอร์มกรอกคำถามทีละข้อ   (~2-3 วัน · ง่ายสุด เกือบเสร็จแล้ว)")
h2("มีอยู่แล้ว (ไม่ต้องแตะ)", color=GREEN)
bullet("ตาราง quiz_patches + quiz_questions + RLS → supabase/migrations/20260521_quiz_patch_system.sql")
bullet("RPC get_active_questions() — เกมดึงใช้แล้ว (QuizManager.FetchFromSupabase, บรรทัด 189)")
bullet("RPC import_quiz_patch() — อัปโหลด JSON ทีเดียว")
bullet("เว็บ admin ทำงานได้ → Assets/StreamingAssets/Web/quiz_admin.html (login/list/toggle/delete/upload/preview/import)")
bullet("เกมดึงจาก Supabase เป็นหลัก (Supabase → cache → local JSON)")
h2("งานที่ต้องทำ")
bullet("SQL: เพิ่ม RPC add_quiz_question(...) แบบ SECURITY DEFINER + GRANT ให้ authenticated", check=True)
bullet("เหตุผล: insert ตรงเข้า quiz_questions ติด RLS (เขียนได้แค่ service_role) — ทำ RPC แบบเดียวกับ import_quiz_patch", level=1)
bullet("validate ใน RPC: choices เป็น array ≥ 2, correct_index 0-3, question ไม่ว่าง", level=1)
bullet("เว็บ: เพิ่มแท็บ “เพิ่มคำถามทีละข้อ” ใน quiz_admin.html", check=True)
bullet("ฟอร์ม: เลือก patch ปลายทาง | คำถาม | 4 ช้อยส์ | radio เลือกข้อถูก | หมวด | ระดับ (easy/medium/hard)", level=1)
bullet("ปุ่ม “บันทึก” → sb.rpc('add_quiz_question', {...}) → toast สำเร็จ → เคลียร์ฟอร์ม", level=1)
bullet("(option) แก้ไข/ลบรายข้อ ในหน้าดูคำถามของ patch (RPC update/delete)", check=True)
bullet("Host เว็บบน GitHub Pages (domain *.supabase.co เปิดเป็นหน้าเว็บไม่ได้ บังคับ text/plain)", check=True)
h2("จุดเสี่ยง", color=RED)
bullet("RLS: ถ้าไม่ทำ RPC จะ insert ไม่ได้ — อย่า insert ตรงด้วย anon key")
bullet("choices ต้องส่งเป็น JSON array จริง ไม่ใช่ string")

# ---------- ข้อ 3 ----------
h1("ข้อ 3 — เก็บ Log การเล่นลง DB   (~2-3 วัน)")
para("หมายเหตุ: “เก็บ log ลง DB” ในโปรเจคจบมักหมายถึง (1) เก็บข้อมูลไปวิเคราะห์ในเล่มวิจัย — น่าจะใช่ที่สุด "
     "(2) ตรวจบั๊กจาก session จริง (3) เป็นหลักฐานตอนสอบ (4) พิสูจน์ว่าระบบ reconnect ทำงาน  "
     "→ ทำ schema แบบ generic ครอบทุกเหตุผล และควรถามอาจารย์ว่าอยากได้ข้อมูลอะไรไปวิเคราะห์", color=MUTED, italic=True)
h2("มีอยู่แล้ว", color=GREEN)
bullet("GameLog.Log — log ทั่วไป ถูก strip ตอน release build (ใช้ debug ใน editor ต่อได้)")
bullet("Pattern เขียน Supabase ชัดเจน: PlayerDataService.CallAuthedFnAsync + REST insert")
h2("SQL: ตาราง game_logs")
code_block([
    "create table public.game_logs (",
    "  id          bigserial primary key,",
    "  user_id     uuid references auth.users(id),",
    "  room_code   text,",
    "  match_id    text,                       -- session name + เวลาเริ่มแมตช์",
    "  event_type  text not null,              -- 'match_start','turn_taken',...",
    "  payload     jsonb default '{}'::jsonb,  -- ข้อมูลเสริมต่อ event",
    "  client_ts   timestamptz,                -- เวลาฝั่งเครื่อง",
    "  created_at  timestamptz not null default timezone('utc', now())",
    ");",
    "alter table public.game_logs enable row level security;",
    "create policy \"insert own logs\" on public.game_logs",
    "  for insert to authenticated with check (user_id = auth.uid());",
    "create index idx_game_logs_match on public.game_logs(match_id);",
    "create index idx_game_logs_event on public.game_logs(event_type);",
])
h2("งานที่ต้องทำ")
bullet("Client: GameLogger (static) — LogEvent(type, payload) → enqueue → flush batch ทุก ~5 วิ/ตอนจบรอบ", check=True)
bullet("แนบ user_id, room_code, match_id อัตโนมัติ; ถ้าออฟไลน์เก็บค้าง flush รอบหน้า (ห้ามค้างเกม)", level=1)
bullet("วาง log point: match_start, turn_taken, card_bought, card_reserved, quiz_answer, "
       "player_disconnect, player_reconnect, match_end", check=True)
bullet("(option) error — hook Application.logMessageReceived เก็บ LogError/Exception", check=True)

# ---------- ข้อ 2 ----------
h1("ข้อ 2 — Tutorial: Coach-mark ทับเกมจริง   (~3-4 วัน)")
h2("มีอยู่แล้ว", color=GREEN)
bullet("TutorialUI.cs = สไลด์โชว์ 3 หน้า (ข้อความ) + TutorialScene + SetupTutorialScene (editor tool)")
bullet("→ เก็บไว้ใช้เป็น “ทฤษฎีก่อนเล่น” ได้ แต่ไม่ใช่ interactive")
h2("งานที่ต้องทำ")
bullet("first-run: PlayerPrefs flag tutorial_done — ครั้งแรกเข้าโหมด tutorial (เล่นกับบอท) อัตโนมัติ", check=True)
bullet("TutorialController + TutorialStep (targetUI, ข้อความสอน, เงื่อนไขผ่าน)", check=True)
bullet("flow ~6 step: หยิบเหรียญ → ซื้อการ์ด → จองการ์ด (กดค้าง) → อธิบายเหรียญดำ → เจอควิซ 1 รอบ → จบ", level=1)
bullet("Coach-mark overlay: Image โปร่งเจาะรูรอบ target + กล่องข้อความ + ลูกศร, บล็อก input นอก target", check=True)
bullet("ผูกเงื่อนไขผ่านกับ action จริงใน GameController (หยิบ/ซื้อ/จอง/ตอบควิซ)", check=True)
bullet("ปุ่ม “ข้าม Tutorial” + ตั้ง tutorial_done=true เมื่อจบ/ข้าม", check=True)
h2("จุดเสี่ยง", color=RED)
bullet("ทำเป็นโหมด bot ล้วน (offline) จะคุมลำดับ step ได้ง่ายกว่า online มาก — แนะนำ")

# ---------- ข้อ 1 ----------
h1("ข้อ 1 — Reconnect: กลับเข้าแมตช์เดิม   (~2-2.5 สัปดาห์ · งานใหญ่สุด)")
h2("มีอยู่แล้ว (~40%)", color=GREEN)
bullet("คนหลุด → seat เป็นบอทเล่นแทน: GameController.Network.cs:133 UpdateDisconnectedPlayerBotStatus")
bullet("late-joiner ขอ full state: RequestFullState()/STATEREQ → host ส่ง board+economy+turn เฉพาะคนนั้น")
bullet("ตอน reconnect มี isBot=false + RequestFullState() แล้ว; seat map เสถียร _seatOrder")
bullet("retry loop ตอน join ห้อง: StartMatchedGameCoroutine (FusionManager.cs:239)")
h2("ช่องว่างหลัก", color=RED)
para("PlayerId เปลี่ยนทุกครั้งที่เข้าห้องใหม่ → _seatOrder ผูก seat กับ PlayerId เก่า (หายไปแล้ว) "
     "คนกลับมากลายเป็น seat ใหม่ ไม่ทวง state เดิม (เหรียญ/การ์ด/คะแนน) ที่บอทถืออยู่", color=RED)
h2("A. Identity ข้าม session (ฝั่ง state) — สัปดาห์ 2")
bullet("เพิ่ม message IDENT|uid (Supabase user id) — ประกาศตอน OnPlayerJoined, broadcast ทั้งห้อง", check=True)
bullet("Authority เก็บ map uid → seat (persistent ตลอดแมตช์, replicate ทุกเครื่อง เผื่อ host migration)", check=True)
bullet("PlayerJoined ใหม่ + รับ IDENT: uid เคยมี seat → “กลับมา” remap _seatOrder[seat]=PlayerId ใหม่, "
       "isBot=false, push full state", check=True)
bullet("uid ใหม่ → seat ใหม่ปกติ", level=1)
h2("B. Orchestration ฝั่ง client — สัปดาห์ 3")
bullet("ดัก OnShutdown/OnDisconnectedFromServer → ถ้า IsGameInProgress → เริ่ม flow reconnect", check=True)
bullet("overlay “กำลังกลับเข้าเกม...” + StartGame(Shared, roomCode เดิม, allowCreate=false) วน retry", check=True)
bullet("หลัง join → ส่ง IDENT → ขอ full state → ปิด overlay", check=True)
bullet("รองรับ app ปิด/เปิดใหม่: roomCode อยู่ใน PlayerPrefs แล้ว → เสนอ “กลับเข้าเกม” ตอนเปิด", check=True)
bullet("เคส host หลุด: คน PlayerId ต่ำสุดถัดไปรับช่วง (มี AuthorityPlayer) — ทดสอบ uid map ยังอยู่", check=True)
bullet("เคสกลับมาตอนเกมจบแล้ว → เด้งหน้าผล/กลับเมนู", check=True)
h2("จุดเสี่ยง (เผื่อเวลาไว้เยอะ)", color=RED)
bullet("Photon session timeout — หลุดนานเกินห้องอาจถูกปิด (ทดสอบ grace window จริง)")
bullet("uid map หายตอน host migration = bug ที่เจอบ่อย → ต้อง replicate ทุกเครื่อง")
bullet("ต้องทดสอบ 2 เครื่องจริง + Android (libnanosockets เคยมีปัญหา)")

# ---------- timeline ----------
h1("ตารางราย 4 สัปดาห์")
table(
    ["สัปดาห์", "งาน"],
    [
        ["สัปดาห์ 1", "ข้อ 4 (ฟอร์มควิซ + RPC + host เว็บ) · ข้อ 3 (ตาราง game_logs + GameLogger + 5-6 event) · ออกแบบ protocol reconnect"],
        ["สัปดาห์ 2", "Reconnect แกนกลาง: IDENT + uid→seat map + seat reclaim + replicate · ทดสอบ ParrelSync 2 instance"],
        ["สัปดาห์ 3", "Reconnect client: auto re-join, overlay, app-restart, เคส host หลุด · ทดสอบ 2 เครื่องจริง + Android"],
        ["สัปดาห์ 4", "ข้อ 2 (TutorialController + coach-mark + first-run) · integration test ทั้ง 4 · ยืนยัน log ลง DB · แก้บั๊กก่อนสอบ"],
    ],
    widths=[26, 156],
)
para("ลำดับนี้กันความเสี่ยง: ถ้า reconnect บานปลาย ยังมี 3 ข้อเสร็จไปสอบได้", italic=True, color=MUTED)

# ---------- checklist ----------
h1("เช็คก่อนเริ่ม (อาจเป็น blocker)")
bullet("มีสิทธิ์รัน SQL migration บน Supabase ไหม (จำเป็นสำหรับข้อ 3 และ 4)", check=True)
bullet("บัญชี GitHub Pages พร้อม host เว็บ admin (ข้อ 4 — เปิดบน supabase.co ไม่ได้)", check=True)
bullet("ParrelSync ติดตั้งแล้วไหม (เทส reconnect 2 จอ)", check=True)
bullet("ถามอาจารย์: อยากได้ข้อมูล log อะไรไปวิเคราะห์ (fine-tune ข้อ 3)", check=True)

out = r"d:\The_Last\GameCardClientUodate\Docs\แผนพัฒนา_4ระบบ_1เดือน.docx"
doc.save(out)
print("SAVED:", out)
